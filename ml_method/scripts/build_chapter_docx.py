"""Assemble the method-chapter drafts into one Word .docx.

Reads the section markdown files in order, strips each file's meta "DRAFT for the thesis..."
paragraph, and renders headings, paragraphs, bold, inline code, bullet lists, and pipe tables into
a docx. The method-architecture section is salvaged from the archived method draft, truncated
before its old "From predictions to a network" section, which the decode section supersedes.

No em-dashes: the drafts contain none, and the script warns if any en-dash or em-dash appears.

Run:  python scripts/build_chapter_docx.py
Out:  docs/gene2net_method_chapter.docx
"""
import os
import re

from docx import Document
from docx.shared import Pt

HERE = os.path.dirname(os.path.abspath(__file__))
DOCS = os.path.join(HERE, "..", "docs")

# (filename relative to docs/, truncate-before-marker or None). Order is the chapter order.
SECTIONS = [
    ("chapter_introduction_draft.md", None),
    ("archive/chapter_method_draft.md", "## From predictions to a network"),
    ("chapter_decode_draft.md", None),
    ("chapter_setup_draft.md", None),
    ("chapter_results_benchmark_draft.md", None),
    ("chapter_diagnostic_draft.md", None),
    ("chapter_retention_draft.md", None),
    ("chapter_feature_importance_draft.md", None),
    ("chapter_limitations_draft.md", None),
    ("chapter_future_work_draft.md", None),
    ("chapter_conclusion_draft.md", None),
]

CHAPTER_TITLE = "A learned detect-then-place method for polyploid network reconstruction"
OUT = os.path.join(DOCS, "gene2net_method_chapter.docx")


def strip_meta(text):
    """Remove the 'DRAFT for the thesis...' meta paragraph (until the next blank line)."""
    out, skip = [], False
    for line in text.split("\n"):
        if line.strip().startswith("DRAFT for the thesis"):
            skip = True
            continue
        if skip:
            if line.strip() == "":
                skip = False
            continue
        out.append(line)
    return "\n".join(out)


def truncate_before(text, marker):
    if marker and marker in text:
        return text[: text.index(marker)]
    return text


def add_runs(paragraph, text):
    """Add text to a paragraph, honoring **bold** and `code` inline markers."""
    pattern = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith("**"):
            paragraph.add_run(tok[2:-2]).bold = True
        else:
            r = paragraph.add_run(tok[1:-1])
            r.font.name = "Consolas"
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def clean_heading(s):
    return s.replace("**", "").replace("`", "").strip()


def parse_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def is_sep(line):
    return re.match(r"^\s*\|?[\s:|-]+\|?[\s:|-]*$", line) is not None and set(line.strip()) <= set("|-: ")


def add_table(doc, rows):
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell_text = row[j] if j < len(row) else ""
            p = table.cell(i, j).paragraphs[0]
            add_runs(p, cell_text)
            if i == 0:
                for run in p.runs:
                    run.bold = True


def render_markdown(doc, text):
    lines = text.split("\n")
    i, para = 0, []

    def flush():
        nonlocal para
        if para:
            add_runs(doc.add_paragraph(), " ".join(para).strip())
            para = []

    while i < len(lines):
        line = lines[i]
        s = line.strip()
        if not s:
            flush(); i += 1; continue
        if s.startswith("### "):
            flush(); doc.add_heading(clean_heading(s[4:]), level=3); i += 1; continue
        if s.startswith("## "):
            flush(); doc.add_heading(clean_heading(s[3:]), level=2); i += 1; continue
        if s.startswith("# "):
            flush(); doc.add_heading(clean_heading(s[2:]), level=1); i += 1; continue
        if s.startswith("|"):
            flush()
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i]); i += 1
            rows = [parse_row(b) for b in block if not is_sep(b)]
            if rows:
                add_table(doc, rows)
            continue
        if re.match(r"^\s*-\s+", line):
            flush()
            while i < len(lines) and re.match(r"^\s*-\s+", lines[i]):
                item = re.sub(r"^\s*-\s+", "", lines[i]).strip()
                add_runs(doc.add_paragraph(style="List Bullet"), item)
                i += 1
            continue
        para.append(s); i += 1
    flush()


def main():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    doc.add_heading(CHAPTER_TITLE, level=0)

    warned = False
    for fname, marker in SECTIONS:
        path = os.path.join(DOCS, fname)
        if not os.path.exists(path):
            raise SystemExit(f"Missing section file: {path}")
        text = open(path, encoding="utf-8").read()
        text = truncate_before(text, marker)
        text = strip_meta(text).strip()
        if ("—" in text or "–" in text) and not warned:
            print(f"WARNING: dash character found in {fname} (em/en dash)")
            warned = True
        render_markdown(doc, text)

    doc.save(OUT)
    print(f"Wrote {OUT}")
    print(f"Sections: {len(SECTIONS)}")


if __name__ == "__main__":
    main()
